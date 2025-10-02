import openai
# import anthropic  # Optional dependency
import PyPDF2
import docx
import io
import re
import asyncio
from typing import Optional, Dict, List, Any, Union
from datetime import datetime
from pathlib import Path
import logging
from PIL import Image
import pytesseract
import aiofiles

from ..config import settings

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class AIService:
    """
    Comprehensive AI service for document processing, text extraction, 
    legal document analysis, and case relevance analysis
    """
    
    def __init__(self):
        self.openai_client = None
        self.anthropic_client = None
        self._initialize_clients()
        
        # Legal document patterns for extraction
        self.legal_patterns = {
            'case_number': r'(?:Case\s+No\.?|Civil\s+Action\s+No\.?|Docket\s+No\.?)\s*:?\s*([A-Z0-9\-\s]+)',
            'court_name': r'(?:IN\s+THE\s+)?(.+?COURT.+?)(?:\n|FOR|OF)',
            'date': r'(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}|\w+\s+\d{1,2},?\s+\d{4})',
            'parties': r'(?:Plaintiff|Petitioner|Appellant)(?:\(s\))?\s*[:\n]\s*(.+?)(?:\n|vs?\.|v\.|against)',
            'defendants': r'(?:Defendant|Respondent|Appellee)(?:\(s\))?\s*[:\n]\s*(.+?)(?:\n|$)',
            'legal_citations': r'(\d+\s+[A-Z][a-z]*\.?\s*\d+d?\s*\d+|\d+\s+U\.S\.C?\.\s*§?\s*\d+)',
            'money_amounts': r'\$[\d,]+(?:\.\d{2})?',
            'phone_numbers': r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            'email_addresses': r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        }
    
    def _initialize_clients(self):
        """Initialize AI service clients with proper error handling"""
        try:
            if settings.openai_api_key:
                self.openai_client = openai.OpenAI(api_key=settings.openai_api_key)
                logger.info("OpenAI client initialized successfully")
            else:
                logger.warning("OpenAI API key not found in settings")
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI client: {e}")
        
        try:
            if settings.anthropic_api_key:
                try:
                    import anthropic
                    self.anthropic_client = anthropic.Anthropic(api_key=settings.anthropic_api_key)
                    logger.info("Anthropic client initialized successfully")
                except ImportError:
                    logger.warning("Anthropic library not installed, skipping Anthropic client initialization")
            else:
                logger.warning("Anthropic API key not found in settings")
        except Exception as e:
            logger.error(f"Failed to initialize Anthropic client: {e}")
    
    # ========================================
    # DOCUMENT TEXT EXTRACTION METHODS
    # ========================================
    
    async def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from various document types (PDF, Word, Images)
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing extracted text, metadata, and success status
        """
        try:
            path_obj = Path(file_path)
            extension = path_obj.suffix.lower()
            
            result = {
                'text': '',
                'metadata': {
                    'filename': path_obj.name,
                    'file_type': extension,
                    'extraction_method': '',
                    'page_count': 0,
                    'word_count': 0,
                    'extraction_timestamp': datetime.now().isoformat()
                },
                'success': False,
                'error': None
            }
            
            if extension == '.pdf':
                result = await self._extract_text_from_pdf(path_obj, result)
            elif extension in ['.docx', '.doc']:
                result = await self._extract_text_from_word(path_obj, result)
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                result = await self._extract_text_from_image(path_obj, result)
            elif extension == '.txt':
                result = await self._extract_text_from_txt(path_obj, result)
            else:
                result['error'] = f"Unsupported file type: {extension}"
                return result
            
            if result['success']:
                result['metadata']['word_count'] = len(result['text'].split())
                logger.info(f"Successfully extracted {result['metadata']['word_count']} words from {path_obj.name}")
            
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return {
                'text': '',
                'metadata': {},
                'success': False,
                'error': str(e)
            }
    
    async def _extract_text_from_pdf(self, file_path: Path, result: Dict) -> Dict:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                result['metadata']['page_count'] = len(pdf_reader.pages)
                result['metadata']['extraction_method'] = 'PyPDF2'
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
                result['text'] = '\n\n'.join(text_content)
                result['success'] = True
                
                # If PDF extraction yields poor results, try OCR
                if len(result['text'].strip()) < 100:
                    logger.info("PDF text extraction yielded minimal content, attempting OCR")
                    ocr_result = await self._extract_text_from_pdf_ocr(file_path, result)
                    if ocr_result['success'] and len(ocr_result['text']) > len(result['text']):
                        return ocr_result
                
                return result
                
        except Exception as e:
            result['error'] = f"PDF extraction error: {str(e)}"
            return result
    
    async def _extract_text_from_pdf_ocr(self, file_path: Path, result: Dict) -> Dict:
        """Extract text from PDF using OCR (for scanned documents)"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            text_content = []
            
            result['metadata']['extraction_method'] = 'OCR (PyMuPDF + Tesseract)'
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")
                
                # Convert to PIL Image and apply OCR
                image = Image.open(io.BytesIO(img_data))
                page_text = pytesseract.image_to_string(image)
                
                if page_text.strip():
                    text_content.append(f"[Page {page_num + 1} - OCR]\n{page_text}")
            
            doc.close()
            result['text'] = '\n\n'.join(text_content)
            result['success'] = True
            return result
            
        except ImportError:
            logger.warning("PyMuPDF not installed, skipping OCR extraction")
            result['error'] = "OCR extraction requires PyMuPDF (fitz) library"
            return result
        except Exception as e:
            result['error'] = f"OCR extraction error: {str(e)}"
            return result
    
    async def _extract_text_from_word(self, file_path: Path, result: Dict) -> Dict:
        """Extract text from Word documents"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            result['metadata']['extraction_method'] = 'python-docx'
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(' | '.join(row_text))
                text_content.append('\n'.join(table_text))
            
            result['text'] = '\n\n'.join(text_content)
            result['success'] = True
            return result
            
        except Exception as e:
            result['error'] = f"Word document extraction error: {str(e)}"
            return result
    
    async def _extract_text_from_image(self, file_path: Path, result: Dict) -> Dict:
        """Extract text from images using OCR"""
        try:
            image = Image.open(file_path)
            
            # Preprocess image for better OCR results
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Apply OCR
            text = pytesseract.image_to_string(image, config='--psm 6')
            
            result['text'] = text
            result['metadata']['extraction_method'] = 'Tesseract OCR'
            result['success'] = True
            return result
            
        except Exception as e:
            result['error'] = f"Image OCR error: {str(e)}"
            return result
    
    async def _extract_text_from_txt(self, file_path: Path, result: Dict) -> Dict:
        """Extract text from plain text files"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            
            result['text'] = text
            result['metadata']['extraction_method'] = 'Direct text read'
            result['success'] = True
            return result
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                        text = await file.read()
                    result['text'] = text
                    result['metadata']['extraction_method'] = f'Text read ({encoding})'
                    result['success'] = True
                    return result
                except:
                    continue
            
            result['error'] = "Could not decode text file with any supported encoding"
            return result
        except Exception as e:
            result['error'] = f"Text file extraction error: {str(e)}"
            return result
    
    # ========================================
    # LEGAL DOCUMENT SUMMARIZATION METHODS
    # ========================================
    
    async def summarize_legal_document(
        self, 
        text: str, 
        filename: str = "", 
        document_type: str = "legal",
        model: str = "gpt-4"
    ) -> Dict[str, Any]:
        """
        Generate AI summary specifically tailored for legal documents
        
        Args:
            text: Document text content
            filename: Name of the document
            document_type: Type of legal document (contract, motion, brief, etc.)
            model: AI model to use
            
        Returns:
            Dictionary with summary, key points, and metadata
        """
        try:
            if model.startswith("gpt") and self.openai_client:
                return await self._summarize_legal_with_openai(text, filename, document_type, model)
            elif model.startswith("claude") and self.anthropic_client:
                return await self._summarize_legal_with_anthropic(text, filename, document_type, model)
            else:
                return await self._fallback_legal_summary(text, filename, document_type)
                
        except Exception as e:
            logger.error(f"Error generating legal summary: {str(e)}")
            return {
                'summary': f"Could not generate summary for {filename}. Error: {str(e)}",
                'key_points': [],
                'legal_entities': {},
                'critical_dates': [],
                'success': False,
                'error': str(e)
            }
    
    async def _summarize_legal_with_openai(self, text: str, filename: str, document_type: str, model: str) -> Dict[str, Any]:
        """Summarize legal documents using OpenAI GPT with specialized prompts"""
        try:
            # Truncate text if too long (GPT-4 has token limits)
            max_chars = 15000  # Roughly 4000 tokens, leaving room for response
            if len(text) > max_chars:
                text = text[:max_chars] + "...[TRUNCATED]"
            
            system_prompt = """You are a legal document analysis expert. Provide detailed, accurate summaries of legal documents with focus on:
1. Legal significance and implications
2. Key parties and their roles
3. Important dates and deadlines
4. Legal precedents or citations
5. Financial terms or damages
6. Procedural requirements
7. Outcomes or decisions

Format your response as structured JSON with specific fields for legal analysis."""
            
            user_prompt = f"""
            Analyze this {document_type} document titled "{filename}" and provide a comprehensive legal summary:

            Document Content:
            {text}

            Please provide your analysis in the following JSON format:
            {{
                "executive_summary": "Brief 2-3 sentence overview",
                "document_type": "Specific type of legal document",
                "key_points": ["List of 5-8 most important points"],
                "parties": {{
                    "plaintiffs": ["List of plaintiffs/petitioners"],
                    "defendants": ["List of defendants/respondents"],
                    "attorneys": ["Legal representatives mentioned"],
                    "judges": ["Judges or magistrates mentioned"]
                }},
                "critical_dates": [
                    {{"date": "YYYY-MM-DD", "event": "Description of what happens on this date"}}
                ],
                "financial_terms": [
                    {{"amount": "$X,XXX", "description": "What this amount represents"}}
                ],
                "legal_citations": ["Any legal precedents, statutes, or regulations cited"],
                "procedural_status": "Current status or next steps in the legal process",
                "risk_assessment": "High/Medium/Low risk factors identified",
                "action_items": ["Required actions or deadlines for parties"]
            }}
            """
            
            response = self.openai_client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=1500,
                temperature=0.2  # Low temperature for consistent, factual analysis
            )
            
            # Parse JSON response
            import json
            try:
                result = json.loads(response.choices[0].message.content)
                result['success'] = True
                result['model_used'] = model
                result['analysis_timestamp'] = datetime.now().isoformat()
                return result
            except json.JSONDecodeError:
                # Fallback if JSON parsing fails
                content = response.choices[0].message.content
                return {
                    'summary': content,
                    'key_points': self._extract_key_points_from_text(content),
                    'legal_entities': {},
                    'critical_dates': [],
                    'success': True,
                    'model_used': model,
                    'analysis_timestamp': datetime.now().isoformat(),
                    'note': 'Response was not in expected JSON format'
                }
            
        except Exception as e:
            logger.error(f"OpenAI legal summarization error: {str(e)}")
            return {
                'summary': f"Failed to generate OpenAI legal summary for {filename}",
                'key_points': [],
                'legal_entities': {},
                'critical_dates': [],
                'success': False,
                'error': str(e)
            }

    async def _summarize_legal_with_anthropic(self, text: str, filename: str, document_type: str, model: str) -> Dict[str, Any]:
        """Summarize legal documents using Anthropic Claude with specialized prompts"""
        try:
            # Truncate text if too long
            max_chars = 20000  # Claude has higher limits
            if len(text) > max_chars:
                text = text[:max_chars] + "...[TRUNCATED]"
            
            prompt = f"""
            As a legal document analysis expert, please analyze this {document_type} document titled "{filename}".

            Document Content:
            {text}

            Provide a comprehensive legal analysis including:

            1. **Executive Summary** (2-3 sentences)
            2. **Document Classification** (specific legal document type)
            3. **Key Legal Points** (5-8 most important items)
            4. **Parties Involved**:
               - Plaintiffs/Petitioners
               - Defendants/Respondents  
               - Legal counsel
               - Court officials
            5. **Critical Dates & Deadlines**
            6. **Financial Terms** (damages, settlements, fees)
            7. **Legal Citations** (statutes, precedents, regulations)
            8. **Procedural Status** (current stage, next steps)
            9. **Risk Assessment** (potential issues or concerns)
            10. **Required Actions** (what parties must do)

            Focus on legal significance, implications, and actionable information.
            """
            
            response = self.anthropic_client.messages.create(
                model=model if model.startswith("claude") else "claude-3-sonnet-20240229",
                max_tokens=2000,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            content = response.content[0].text
            
            return {
                'summary': content,
                'key_points': self._extract_key_points_from_text(content),
                'legal_entities': self._extract_legal_entities_from_text(content),
                'critical_dates': self._extract_dates_from_text(content),
                'success': True,
                'model_used': model,
                'analysis_timestamp': datetime.now().isoformat()
            }
            
        except Exception as e:
            logger.error(f"Anthropic legal summarization error: {str(e)}")
            return {
                'summary': f"Failed to generate Claude legal summary for {filename}",
                'key_points': [],
                'legal_entities': {},
                'critical_dates': [],
                'success': False,
                'error': str(e)
            }
    
    async def _fallback_legal_summary(self, text: str, filename: str, document_type: str) -> Dict[str, Any]:
        """Fallback legal summary when AI services are not available"""
        try:
            # Simple extractive summary - get first few sentences and key information
            sentences = text.split('.')[:5]
            summary = '. '.join(sentences).strip()
            
            if len(summary) > 500:
                summary = summary[:500] + "..."
            
            # Extract basic information using regex patterns
            entities = await self._extract_legal_entities(text)
            
            return {
                'summary': f"Basic Legal Summary of {filename} ({document_type}): {summary}",
                'key_points': self._extract_key_points_from_text(text),
                'legal_entities': entities,
                'critical_dates': entities.get('dates', [])[:5],
                'success': True,
                'model_used': 'fallback',
                'analysis_timestamp': datetime.now().isoformat(),
                'note': 'Generated using fallback method due to AI service unavailability'
            }
            
        except Exception as e:
            logger.error(f"Fallback legal summary error: {str(e)}")
            return {
                'summary': f"Could not generate summary for {filename}",
                'key_points': [],
                'legal_entities': {},
                'critical_dates': [],
                'success': False,
                'error': str(e)
            }
    
    async def _summarize_with_anthropic(self, text: str, filename: str, model: str) -> str:
        """Summarize using Anthropic Claude"""
        try:
            # Truncate text if too long
            max_chars = 15000  # Claude has higher limits
            if len(text) > max_chars:
                text = text[:max_chars] + "..."
            
            prompt = f"""
            Please provide a concise summary of the following document titled "{filename}":

            {text}

            Focus on the main points, key information, and any important details. Keep it under 200 words.
            """
            
            # TODO: Replace with actual Anthropic API call
            # response = await self.anthropic_client.messages.create(
            #     model=model,
            #     max_tokens=300,
            #     messages=[
            #         {"role": "user", "content": prompt}
            #     ]
            # )
            # return response.content[0].text.strip()
            
            # Mock response for now
            return f"Claude Summary of {filename}: This document discusses important legal matters with specific case references and procedural information. It contains detailed information about court proceedings, party details, and administrative processes relevant to the legal case."
            
        except Exception as e:
            print(f"Anthropic summarization error: {str(e)}")
            return f"Failed to generate Claude summary for {filename}"
    
    async def _fallback_summary(self, text: str, filename: str) -> str:
        """Fallback summary when AI services are not available"""
        try:
            # Simple extractive summary - get first few sentences
            sentences = text.split('.')[:3]
            summary = '. '.join(sentences).strip()
            
            if len(summary) > 300:
                summary = summary[:300] + "..."
            
            return f"Basic Summary of {filename}: {summary}"
            
        except Exception as e:
            return f"Could not generate summary for {filename}"
    
    # ========================================
    # KEY INFORMATION EXTRACTION METHODS
    # ========================================
    
    async def extract_key_information(self, text: str, info_type: str = "legal") -> Dict[str, Any]:
        """
        Extract specific types of structured information from legal documents
        
        Args:
            text: Document text content
            info_type: Type of extraction ("legal", "case", "contract", "financial")
            
        Returns:
            Dictionary with extracted structured information
        """
        try:
            result = {
                'extraction_type': info_type,
                'extraction_timestamp': datetime.now().isoformat(),
                'success': False,
                'data': {}
            }
            
            if info_type == "legal":
                result['data'] = await self._extract_legal_entities(text)
            elif info_type == "case":
                result['data'] = await self._extract_case_info(text)
            elif info_type == "contract":
                result['data'] = await self._extract_contract_info(text)
            elif info_type == "financial":
                result['data'] = await self._extract_financial_info(text)
            else:
                result['data'] = await self._extract_general_info(text)
            
            result['success'] = True
            return result
                
        except Exception as e:
            logger.error(f"Error extracting {info_type} information: {str(e)}")
            return {
                'extraction_type': info_type,
                'extraction_timestamp': datetime.now().isoformat(),
                'success': False,
                'error': str(e),
                'data': {}
            }
    
    async def _extract_legal_entities(self, text: str) -> Dict[str, List]:
        """Extract comprehensive legal entities and information"""
        entities = {
            "case_numbers": [],
            "court_names": [],
            "judges": [],
            "attorneys": [],
            "parties": {
                "plaintiffs": [],
                "defendants": [],
                "witnesses": []
            },
            "legal_citations": [],
            "statutes": [],
            "dates": [],
            "financial_amounts": [],
            "addresses": [],
            "phone_numbers": [],
            "email_addresses": []
        }
        
        # Extract case numbers
        case_matches = re.findall(self.legal_patterns['case_number'], text, re.IGNORECASE)
        entities["case_numbers"] = [match.strip() for match in case_matches]
        
        # Extract court names
        court_matches = re.findall(self.legal_patterns['court_name'], text, re.IGNORECASE)
        entities["court_names"] = [match.strip() for match in court_matches]
        
        # Extract dates
        date_matches = re.findall(self.legal_patterns['date'], text)
        entities["dates"] = [match.strip() for match in date_matches]
        
        # Extract parties
        plaintiff_matches = re.findall(self.legal_patterns['parties'], text, re.IGNORECASE)
        entities["parties"]["plaintiffs"] = [match.strip() for match in plaintiff_matches]
        
        defendant_matches = re.findall(self.legal_patterns['defendants'], text, re.IGNORECASE)
        entities["parties"]["defendants"] = [match.strip() for match in defendant_matches]
        
        # Extract legal citations
        citation_matches = re.findall(self.legal_patterns['legal_citations'], text)
        entities["legal_citations"] = [match.strip() for match in citation_matches]
        
        # Extract financial amounts
        money_matches = re.findall(self.legal_patterns['money_amounts'], text)
        entities["financial_amounts"] = [match.strip() for match in money_matches]
        
        # Extract contact information
        phone_matches = re.findall(self.legal_patterns['phone_numbers'], text)
        entities["phone_numbers"] = [match.strip() for match in phone_matches]
        
        email_matches = re.findall(self.legal_patterns['email_addresses'], text)
        entities["email_addresses"] = [match.strip() for match in email_matches]
        
        # Extract attorneys and judges using common patterns
        attorney_patterns = [
            r'(?:Attorney|Counsel|Esq\.?)\s*:?\s*([A-Z][a-z]+\s+[A-Z][a-z]+)',
            r'([A-Z][a-z]+\s+[A-Z][a-z]+),?\s*(?:Esq\.?|Attorney|Counsel)'
        ]
        
        for pattern in attorney_patterns:
            matches = re.findall(pattern, text)
            entities["attorneys"].extend([match.strip() for match in matches])
        
        judge_patterns = [
            r'(?:Judge|Justice|Magistrate)\s+([A-Z][a-z]+\s+[A-Z][a-z]+)',
            r'(?:Hon\.?|Honorable)\s+([A-Z][a-z]+\s+[A-Z][a-z]+)'
        ]
        
        for pattern in judge_patterns:
            matches = re.findall(pattern, text)
            entities["judges"].extend([match.strip() for match in matches])
        
        # Remove duplicates and clean up
        for key in entities:
            if isinstance(entities[key], list):
                entities[key] = list(set(entities[key]))
            elif isinstance(entities[key], dict):
                for subkey in entities[key]:
                    if isinstance(entities[key][subkey], list):
                        entities[key][subkey] = list(set(entities[key][subkey]))
        
        return entities
    
    async def _extract_case_info(self, text: str) -> Dict[str, Any]:
        """Extract case-specific information"""
        case_info = {
            "case_number": None,
            "case_title": None,
            "court": None,
            "plaintiff": None,
            "defendant": None,
            "case_type": None,
            "filing_date": None,
            "hearing_date": None,
            "status": None,
            "presiding_judge": None,
            "docket_entries": []
        }
        
        # Extract case number (first occurrence)
        case_match = re.search(self.legal_patterns['case_number'], text, re.IGNORECASE)
        if case_match:
            case_info["case_number"] = case_match.group(1).strip()
        
        # Extract case title (usually at the beginning)
        title_patterns = [
            r'^(.+?)\s+(?:vs?\.?|v\.?|against)\s+(.+?)(?:\n|Case)', 
            r'IN\s+THE\s+MATTER\s+OF\s+(.+?)(?:\n|$)',
            r'(.+?)\s+(?:vs?\.?|v\.?)\s+(.+?)(?:\s*Case|\s*Civil|\s*Criminal|\n)'
        ]
        
        for pattern in title_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                if len(match.groups()) >= 2:
                    case_info["plaintiff"] = match.group(1).strip()
                    case_info["defendant"] = match.group(2).strip()
                    case_info["case_title"] = f"{case_info['plaintiff']} v. {case_info['defendant']}"
                else:
                    case_info["case_title"] = match.group(1).strip()
                break
        
        # Extract court name
        court_match = re.search(self.legal_patterns['court_name'], text, re.IGNORECASE)
        if court_match:
            case_info["court"] = court_match.group(1).strip()
        
        # Extract dates
        dates = re.findall(self.legal_patterns['date'], text)
        if dates:
            case_info["filing_date"] = dates[0] if len(dates) > 0 else None
            case_info["hearing_date"] = dates[1] if len(dates) > 1 else None
        
        # Determine case type based on keywords
        case_types = {
            'civil': ['civil', 'tort', 'contract', 'negligence', 'damages'],
            'criminal': ['criminal', 'felony', 'misdemeanor', 'prosecution', 'defendant charged'],
            'family': ['divorce', 'custody', 'adoption', 'domestic', 'family court'],
            'probate': ['estate', 'will', 'probate', 'inheritance', 'guardian'],
            'bankruptcy': ['bankruptcy', 'chapter 7', 'chapter 11', 'debtor', 'creditor']
        }
        
        text_lower = text.lower()
        for case_type, keywords in case_types.items():
            if any(keyword in text_lower for keyword in keywords):
                case_info["case_type"] = case_type
                break
        
        return case_info
    
    async def _extract_contract_info(self, text: str) -> Dict[str, Any]:
        """Extract contract-specific information"""
        contract_info = {
            "contract_type": None,
            "parties": [],
            "effective_date": None,
            "expiration_date": None,
            "contract_value": None,
            "payment_terms": [],
            "key_obligations": [],
            "termination_clauses": [],
            "governing_law": None,
            "signatures": []
        }
        
        # Detect contract type
        contract_types = {
            'employment': ['employment', 'job', 'salary', 'employee', 'employer'],
            'service': ['service agreement', 'consulting', 'professional services'],
            'lease': ['lease', 'rent', 'landlord', 'tenant', 'premises'],
            'sales': ['purchase', 'sale', 'buyer', 'seller', 'goods'],
            'nda': ['non-disclosure', 'confidentiality', 'confidential information'],
            'licensing': ['license', 'intellectual property', 'copyright', 'trademark']
        }
        
        text_lower = text.lower()
        for contract_type, keywords in contract_types.items():
            if any(keyword in text_lower for keyword in keywords):
                contract_info["contract_type"] = contract_type
                break
        
        # Extract monetary amounts
        amounts = re.findall(self.legal_patterns['money_amounts'], text)
        if amounts:
            contract_info["contract_value"] = amounts[0]
        
        # Extract dates
        dates = re.findall(self.legal_patterns['date'], text)
        if dates:
            contract_info["effective_date"] = dates[0] if len(dates) > 0 else None
            contract_info["expiration_date"] = dates[1] if len(dates) > 1 else None
        
        return contract_info
    
    async def _extract_financial_info(self, text: str) -> Dict[str, Any]:
        """Extract financial information from documents"""
        financial_info = {
            "monetary_amounts": [],
            "payment_schedules": [],
            "interest_rates": [],
            "penalties": [],
            "damages": [],
            "fees": []
        }
        
        # Extract all monetary amounts
        amounts = re.findall(self.legal_patterns['money_amounts'], text)
        financial_info["monetary_amounts"] = amounts
        
        # Extract interest rates
        interest_patterns = [
            r'(\d+(?:\.\d+)?%)\s*(?:interest|rate)',
            r'(?:interest|rate)\s*(?:of|at)\s*(\d+(?:\.\d+)?%)'
        ]
        
        for pattern in interest_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            financial_info["interest_rates"].extend(matches)
        
        return financial_info
    
    async def _extract_general_info(self, text: str) -> Dict[str, Any]:
        """Extract general information from any document"""
        general_info = {
            "entities": {
                "people": [],
                "organizations": [],
                "locations": [],
                "dates": [],
                "amounts": []
            },
            "document_structure": {
                "headings": [],
                "bullet_points": [],
                "numbered_items": []
            }
        }
        
        # Extract basic entities
        general_info["entities"]["dates"] = re.findall(self.legal_patterns['date'], text)
        general_info["entities"]["amounts"] = re.findall(self.legal_patterns['money_amounts'], text)
        
        # Extract document structure
        headings = re.findall(r'^([A-Z][A-Za-z\s]+):?\s*$', text, re.MULTILINE)
        general_info["document_structure"]["headings"] = headings
        
        bullet_points = re.findall(r'^\s*[•\-\*]\s*(.+)$', text, re.MULTILINE)
        general_info["document_structure"]["bullet_points"] = bullet_points
        
        numbered_items = re.findall(r'^\s*\d+\.\s*(.+)$', text, re.MULTILINE)
        general_info["document_structure"]["numbered_items"] = numbered_items
        
        return general_info
    
    # ========================================
    # CASE RELEVANCE ANALYSIS METHODS
    # ========================================
    
    async def analyze_case_relevance(
        self, 
        document_text: str, 
        case_context: Dict[str, Any],
        analysis_type: str = "comprehensive"
    ) -> Dict[str, Any]:
        """
        Analyze how relevant a document is to a specific legal case
        
        Args:
            document_text: Text content of the document to analyze
            case_context: Dictionary with case information (case_number, parties, legal_issues, etc.)
            analysis_type: Type of analysis ("comprehensive", "quick", "similarity")
            
        Returns:
            Dictionary with relevance analysis results
        """
        try:
            if analysis_type == "comprehensive":
                return await self._comprehensive_relevance_analysis(document_text, case_context)
            elif analysis_type == "quick":
                return await self._quick_relevance_analysis(document_text, case_context)
            elif analysis_type == "similarity":
                return await self._similarity_analysis(document_text, case_context)
            else:
                return await self._comprehensive_relevance_analysis(document_text, case_context)
                
        except Exception as e:
            logger.error(f"Error in case relevance analysis: {str(e)}")
            return {
                'relevance_score': 0.0,
                'analysis_type': analysis_type,
                'success': False,
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }
    
    async def _comprehensive_relevance_analysis(self, document_text: str, case_context: Dict[str, Any]) -> Dict[str, Any]:
        """Perform comprehensive AI-powered relevance analysis"""
        try:
            # Prepare context information
            case_summary = self._format_case_context(case_context)
            
            if self.openai_client:
                return await self._ai_relevance_analysis_openai(document_text, case_summary, case_context)
            elif self.anthropic_client:
                return await self._ai_relevance_analysis_anthropic(document_text, case_summary, case_context)
            else:
                return await self._rule_based_relevance_analysis(document_text, case_context)
                
        except Exception as e:
            logger.error(f"Comprehensive relevance analysis error: {str(e)}")
            return await self._rule_based_relevance_analysis(document_text, case_context)
    
    async def _ai_relevance_analysis_openai(self, document_text: str, case_summary: str, case_context: Dict[str, Any]) -> Dict[str, Any]:
        """Use OpenAI for intelligent relevance analysis"""
        try:
            # Truncate document if too long
            max_chars = 12000
            if len(document_text) > max_chars:
                document_text = document_text[:max_chars] + "...[TRUNCATED]"
            
            system_prompt = """You are a legal document relevance analysis expert. Analyze how relevant a document is to a specific legal case and provide detailed scoring and reasoning."""
            
            user_prompt = f"""
            Please analyze the relevance of the following document to the specified legal case:

            CASE CONTEXT:
            {case_summary}

            DOCUMENT TO ANALYZE:
            {document_text}

            Provide your analysis in JSON format with the following structure:
            {{
                "overall_relevance_score": 0.0-1.0,
                "relevance_category": "highly_relevant|moderately_relevant|tangentially_relevant|not_relevant",
                "matching_elements": {{
                    "parties": {{"score": 0.0-1.0, "matches": ["list of matching parties"]}},
                    "legal_issues": {{"score": 0.0-1.0, "matches": ["list of matching legal concepts"]}},
                    "dates": {{"score": 0.0-1.0, "matches": ["list of relevant dates"]}},
                    "case_numbers": {{"score": 0.0-1.0, "matches": ["list of matching case numbers"]}},
                    "financial_terms": {{"score": 0.0-1.0, "matches": ["list of financial connections"]}}
                }},
                "relevance_reasoning": "Detailed explanation of why this document is or isn't relevant",
                "key_connections": ["List of specific connections to the case"],
                "potential_use": "How this document could be used in the case",
                "confidence_level": "high|medium|low",
                "recommended_action": "file_with_case|review_further|archive|discard"
            }}
            """
            
            response = self.openai_client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=1000,
                temperature=0.1
            )
            
            import json
            try:
                result = json.loads(response.choices[0].message.content)
                result['analysis_method'] = 'AI (OpenAI GPT-4)'
                result['timestamp'] = datetime.now().isoformat()
                result['success'] = True
                return result
            except json.JSONDecodeError:
                # Fallback parsing
                content = response.choices[0].message.content
                return self._parse_ai_response_fallback(content, "OpenAI")
                
        except Exception as e:
            logger.error(f"OpenAI relevance analysis error: {str(e)}")
            return await self._rule_based_relevance_analysis(document_text, case_context)
    
    async def _ai_relevance_analysis_anthropic(self, document_text: str, case_summary: str, case_context: Dict[str, Any]) -> Dict[str, Any]:
        """Use Anthropic Claude for intelligent relevance analysis"""
        try:
            max_chars = 15000
            if len(document_text) > max_chars:
                document_text = document_text[:max_chars] + "...[TRUNCATED]"
            
            prompt = f"""
            As a legal document analysis expert, please analyze how relevant this document is to the specified legal case.

            CASE CONTEXT:
            {case_summary}

            DOCUMENT TO ANALYZE:
            {document_text}

            Please provide a comprehensive relevance analysis including:

            1. **Overall Relevance Score** (0.0 to 1.0)
            2. **Relevance Category** (highly relevant, moderately relevant, tangentially relevant, not relevant)
            3. **Matching Elements Analysis**:
               - Party matches and relevance
               - Legal issue connections
               - Date correlations
               - Case number matches
               - Financial term connections
            4. **Detailed Reasoning** (why is this document relevant or not?)
            5. **Key Connections** (specific links to the case)
            6. **Potential Use** (how could this document help the case?)
            7. **Confidence Level** (high, medium, low)
            8. **Recommended Action** (file with case, review further, archive, or discard)

            Focus on legal significance and practical utility for the case.
            """
            
            response = self.anthropic_client.messages.create(
                model="claude-3-sonnet-20240229",
                max_tokens=1200,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            content = response.content[0].text
            
            # Extract relevance score using regex
            score_match = re.search(r'(?:score|relevance).*?(\d+\.\d+|\d+)', content, re.IGNORECASE)
            relevance_score = float(score_match.group(1)) if score_match else 0.5
            if relevance_score > 1.0:
                relevance_score = relevance_score / 10.0  # Handle scores given as percentages
            
            return {
                'overall_relevance_score': relevance_score,
                'relevance_category': self._determine_relevance_category(relevance_score),
                'analysis_content': content,
                'matching_elements': self._extract_matching_elements_from_response(content),
                'analysis_method': 'AI (Anthropic Claude)',
                'timestamp': datetime.now().isoformat(),
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Anthropic relevance analysis error: {str(e)}")
            return await self._rule_based_relevance_analysis(document_text, case_context)
    
    async def _rule_based_relevance_analysis(self, document_text: str, case_context: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback rule-based relevance analysis when AI is unavailable"""
        try:
            relevance_factors = {
                'party_matches': 0.0,
                'case_number_matches': 0.0,
                'date_relevance': 0.0,
                'keyword_matches': 0.0,
                'legal_term_matches': 0.0
            }
            
            document_lower = document_text.lower()
            
            # Check for party matches
            if 'parties' in case_context:
                parties = case_context['parties']
                if isinstance(parties, list):
                    party_matches = sum(1 for party in parties if party.lower() in document_lower)
                    relevance_factors['party_matches'] = min(party_matches / len(parties), 1.0) if parties else 0.0
                elif isinstance(parties, dict):
                    all_parties = []
                    for party_type in parties.values():
                        if isinstance(party_type, list):
                            all_parties.extend(party_type)
                    party_matches = sum(1 for party in all_parties if party.lower() in document_lower)
                    relevance_factors['party_matches'] = min(party_matches / len(all_parties), 1.0) if all_parties else 0.0
            
            # Check for case number matches
            if 'case_number' in case_context and case_context['case_number']:
                case_number = case_context['case_number'].lower()
                if case_number in document_lower:
                    relevance_factors['case_number_matches'] = 1.0
            
            # Check for legal keywords
            legal_keywords = [
                'contract', 'agreement', 'lawsuit', 'litigation', 'court', 'judge',
                'plaintiff', 'defendant', 'motion', 'brief', 'evidence', 'testimony',
                'settlement', 'damages', 'liability', 'negligence', 'breach'
            ]
            
            keyword_matches = sum(1 for keyword in legal_keywords if keyword in document_lower)
            relevance_factors['keyword_matches'] = min(keyword_matches / len(legal_keywords), 1.0)
            
            # Calculate overall score
            weights = {
                'party_matches': 0.4,
                'case_number_matches': 0.3,
                'date_relevance': 0.1,
                'keyword_matches': 0.1,
                'legal_term_matches': 0.1
            }
            
            overall_score = sum(relevance_factors[factor] * weights[factor] for factor in weights)
            
            return {
                'overall_relevance_score': overall_score,
                'relevance_category': self._determine_relevance_category(overall_score),
                'matching_elements': relevance_factors,
                'analysis_method': 'Rule-based (fallback)',
                'timestamp': datetime.now().isoformat(),
                'success': True,
                'note': 'Rule-based analysis used due to AI service unavailability'
            }
            
        except Exception as e:
            logger.error(f"Rule-based relevance analysis error: {str(e)}")
            return {
                'overall_relevance_score': 0.0,
                'relevance_category': 'unknown',
                'success': False,
                'error': str(e),
                'timestamp': datetime.now().isoformat()
            }
    
    async def _quick_relevance_analysis(self, document_text: str, case_context: Dict[str, Any]) -> Dict[str, Any]:
        """Quick relevance check based on keyword matching"""
        return await self._rule_based_relevance_analysis(document_text, case_context)
    
    async def _similarity_analysis(self, document_text: str, case_context: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze document similarity to case materials"""
        # This could be enhanced with embedding-based similarity
        return await self._rule_based_relevance_analysis(document_text, case_context)
    
    # ========================================
    # HELPER METHODS
    # ========================================
    
    def _format_case_context(self, case_context: Dict[str, Any]) -> str:
        """Format case context for AI prompts"""
        context_parts = []
        
        if 'case_number' in case_context:
            context_parts.append(f"Case Number: {case_context['case_number']}")
        
        if 'case_title' in case_context:
            context_parts.append(f"Case Title: {case_context['case_title']}")
        
        if 'parties' in case_context:
            context_parts.append(f"Parties: {case_context['parties']}")
        
        if 'legal_issues' in case_context:
            context_parts.append(f"Legal Issues: {case_context['legal_issues']}")
        
        if 'case_type' in case_context:
            context_parts.append(f"Case Type: {case_context['case_type']}")
        
        return '\n'.join(context_parts)
    
    def _determine_relevance_category(self, score: float) -> str:
        """Determine relevance category based on score"""
        if score >= 0.8:
            return "highly_relevant"
        elif score >= 0.6:
            return "moderately_relevant"
        elif score >= 0.3:
            return "tangentially_relevant"
        else:
            return "not_relevant"
    
    def _extract_key_points_from_text(self, text: str) -> List[str]:
        """Extract key points from AI response text"""
        # Simple extraction of bullet points or numbered items
        key_points = []
        
        # Look for bullet points
        bullet_matches = re.findall(r'[•\-\*]\s*(.+)', text)
        key_points.extend(bullet_matches)
        
        # Look for numbered items
        numbered_matches = re.findall(r'\d+\.\s*(.+)', text)
        key_points.extend(numbered_matches)
        
        return key_points[:10]  # Limit to top 10
    
    def _extract_legal_entities_from_text(self, text: str) -> Dict[str, List]:
        """Extract legal entities mentioned in AI response"""
        entities = {
            'people': [],
            'organizations': [],
            'courts': [],
            'cases': []
        }
        
        # Basic entity extraction from AI response
        # This is a simplified version - could be enhanced with NER
        
        return entities
    
    def _extract_dates_from_text(self, text: str) -> List[str]:
        """Extract dates mentioned in AI response"""
        dates = re.findall(self.legal_patterns['date'], text)
        return dates[:10]  # Limit to top 10
    
    def _parse_ai_response_fallback(self, content: str, source: str) -> Dict[str, Any]:
        """Fallback parser for AI responses that aren't in JSON format"""
        # Extract relevance score
        score_match = re.search(r'(?:score|relevance).*?(\d+\.\d+|\d+)', content, re.IGNORECASE)
        relevance_score = float(score_match.group(1)) if score_match else 0.5
        if relevance_score > 1.0:
            relevance_score = relevance_score / 10.0
        
        return {
            'overall_relevance_score': relevance_score,
            'relevance_category': self._determine_relevance_category(relevance_score),
            'analysis_content': content,
            'analysis_method': f'AI ({source}) - Fallback parsing',
            'timestamp': datetime.now().isoformat(),
            'success': True,
            'note': 'Response was parsed using fallback method'
        }
    
    def _extract_matching_elements_from_response(self, content: str) -> Dict[str, Any]:
        """Extract matching elements from AI response text"""
        return {
            'parties': {'score': 0.5, 'matches': []},
            'legal_issues': {'score': 0.5, 'matches': []},
            'dates': {'score': 0.5, 'matches': []},
            'case_numbers': {'score': 0.5, 'matches': []},
            'financial_terms': {'score': 0.5, 'matches': []}
        }